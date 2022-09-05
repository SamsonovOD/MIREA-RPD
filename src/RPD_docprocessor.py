import re
from docx import Document
import pprint

# Импорт файла
def docload(path, file):
    # Если расширение файла не .docx, пропустить
    if re.compile('.*.docx$').match(file) != None:
        return Document(path+"\\"+file)
    else:
        return False

# Убрать пробелы
def trim_cell_spaces(cell):
    text = cell.text
    while text[0] == " ":
        text = text[1:]
    while text[-1] == " ":
        text = text[:-2]
    return text

# Изъятие текста из документа   
def parse_doc_topics(doc):
    f_in_word_area = False
    topics_list = []
    for table in doc.tables:
        for cell in table.column_cells(2):
            # Предположить что темы всегда находятся во второй колонке и мы знаем заначения ячеек в строках до и после интересующей области
            if f_in_word_area:
                topics_list.append(cell.text)
            if "Наименование разделов и тем" in cell.text:
                f_in_word_area = True
            if "ОЦЕНОЧНЫЕ МАТЕРИАЛЫ" in cell.text:
                f_in_word_area = False
    return topics_list
    
def parse_doc_topics(doc, rpd_bigdata):
    f_in_word_area = False
    topics_list = []
    speciality_id = 0
    var_module = ""
    var_speciality_code = ""
    var_speciality_name = ""
    var_direction = ""
    var_department = ""
    path_list = []
    for table in doc.tables:
        cap_intro = False
        cap_department = False
        cap_speciality = False
        cap_direction = False
        cap_module = False
        cap_lectures = False
        for row in table.rows:
            for cell in row.cells:
                if "Рабочая программа дисциплины (модуля)" in cell.text:
                    cap_module = True
                    cap_intro = True
                    path_list.clear()
                if cap_module and cell.text != "" and not "Рабочая программа дисциплины (модуля)" in cell.text:
                    var_module = trim_cell_spaces(cell)
                    cap_module = False
                    
                if "Читающее подразделение" in cell.text and cap_intro:
                    cap_department = True
                if cap_department and cell.text != "" and not "Читающее подразделение" in cell.text:
                    var_department = trim_cell_spaces(cell)
                    cap_department = False

                if "Направление" in cell.text and cap_intro:
                    cap_speciality = True
                if cap_speciality and cell.text != "" and not "Направление" in cell.text:
                    var_text = trim_cell_spaces(cell)
                    var_code = var_text.split(" ")[0]
                    var_speciality_code = var_code
                    var_speciality_name = var_text.replace(var_code+" ", "")
                    cap_speciality = False
                    
                if "Направленность" in cell.text and cap_intro:
                    cap_direction = True
                if cap_direction and cell.text != "" and not "Направленность" in cell.text:
                    var_direction = trim_cell_spaces(cell)
                    cap_direction = False
                    
                if "Форма обучения" in cell.text and cap_intro:
                    cap_intro = False
                    path_list.append([var_speciality_name, var_direction, var_department, var_module])
                    path_list.append([{"code": var_speciality_code}, {}, {}, {}])
                    rpd_bigdata.setdefault(path_list[0][0], path_list[1][0]).setdefault(path_list[0][1], path_list[1][1]).setdefault(path_list[0][2], path_list[1][2]).setdefault(path_list[0][3], path_list[1][3])
                    rpd_bigdata[path_list[0][0]][path_list[0][1]][path_list[0][2]][path_list[0][3]] = []
                    
        for cell in table.column_cells(2):
            if "Наименование разделов и тем" in cell.text:
                cap_lectures = True
            if cap_lectures:
                rpd_bigdata[path_list[0][0]][path_list[0][1]][path_list[0][2]][path_list[0][3]].append(trim_cell_spaces(cell))
            if "ОЦЕНОЧНЫЕ МАТЕРИАЛЫ" in cell.text:
                cap_lectures = False
    return rpd_bigdata

# Удалить элемент из текстового массива
def remove_filter_item(text_list, remove_match):
    regex = re.compile(r''+remove_match)
    text_list = [i for i in text_list if not regex.match(i)]
    return text_list

# Пред-обработка списка тем
def clean_topics(topics_list):
    # Убрать пустые ячейки
    topics_list = remove_filter_item(topics_list, '^\s*$')
    # Убрать ячейки с нумерацей разделов
    topics_list = remove_filter_item(topics_list, '\s\d\.\s.*$') 
    # Убрать ячеки с практическими занятиям
    topics_list = remove_filter_item(topics_list, '.*(Пр).*')
    return topics_list
    
# Представить имя файла как объект
def group_doc_name(doc_name):
    doc_param = {}
    try:
        doc_param["year"] = re.search('\d+-\d+', doc_name).group(0)
        doc_param["code"] = re.search('_\d\d_\d\d_\d\d', doc_name).group(0)[1:]
        doc_param["dept"] = re.search('[А-Я]+_[А-Я]+', doc_name).group(0)
        doc_param["subject"] = doc_name.split("_plx_")[1].split(".doc")[0]
    except:
        doc_param["year"] = "2022"
        doc_param["code"] = "09.03.03"
        doc_param["dept"] = "ИИТ"
        doc_param["subject"] = "НДисциплина"
    return doc_param
